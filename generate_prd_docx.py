from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.shared import OxmlElement as SharedOxmlElement


OUTPUT = "E:/炖时光/炖时光微信小程序PRD_v1.1.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="D9DEE5", size="8"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_table_width(table, width_inches):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def set_col_width(cell, inches):
    cell.width = Inches(inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=11, bold=False, color="000000", name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(p, before=0, after=6, line=1.25, keep_with_next=False):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_with_next = keep_with_next


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 1:
        r = p.add_run(text)
        set_run_font(r, size=18, bold=True, color="1F1F1F")
        format_paragraph(p, before=18, after=8, line=1.1, keep_with_next=True)
    elif level == 2:
        r = p.add_run(text)
        set_run_font(r, size=14, bold=True, color="1F1F1F")
        format_paragraph(p, before=14, after=6, line=1.1, keep_with_next=True)
    else:
        r = p.add_run(text)
        set_run_font(r, size=12, bold=True, color="43505A")
        format_paragraph(p, before=10, after=4, line=1.1, keep_with_next=True)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=11, bold=False)
    format_paragraph(p, before=0, after=6, line=1.25)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    if level == 1:
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
    r = p.add_run(text)
    set_run_font(r, size=11)
    format_paragraph(p, before=0, after=4, line=1.2)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r, size=11)
    format_paragraph(p, before=0, after=4, line=1.2)
    return p


def add_label_value_table(doc, rows, widths=(1.7, 4.8)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 6.5)
    set_table_borders(table)
    for label, value in rows:
        row = table.add_row()
        row.height = Pt(24)
        c1, c2 = row.cells
        c1.text = ""
        c2.text = ""
        p1 = c1.paragraphs[0]
        p2 = c2.paragraphs[0]
        r1 = p1.add_run(label)
        r2 = p2.add_run(value)
        set_run_font(r1, size=10.5, bold=True, color="1F1F1F")
        set_run_font(r2, size=10.5, color="1F1F1F")
        format_paragraph(p1, before=0, after=0, line=1.0)
        format_paragraph(p2, before=0, after=0, line=1.0)
        c1.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        c2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(c1)
        set_cell_margins(c2)
        set_col_width(c1, widths[0])
        set_col_width(c2, widths[1])
        set_cell_shading(c1, "F2F4F7")
    return table


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_width(table, 6.5)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_font(r, size=10.5, bold=True, color="1F1F1F")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p, before=0, after=0, line=1.0)
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        set_col_width(cell, widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row_data in rows:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run_font(r, size=10.2, color="1F1F1F")
            format_paragraph(p, before=0, after=0, line=1.05)
            set_cell_margins(cell)
            set_col_width(cell, widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    return table


def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("炖时光微信小程序 PRD v1.1")
    set_run_font(r, size=9, color="666666")


def main():
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_page_number_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    for sname, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        s = styles[sname]
        s.font.name = "Microsoft YaHei"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size)

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("炖时光微信小程序 PRD v1.1")
    set_run_font(r, size=22, bold=True, color="111111")
    format_paragraph(p, before=0, after=4, line=1.0)

    p = doc.add_paragraph()
    r = p.add_run("适用于门店点单、预订备货、会员复购、储值活动与经营统计")
    set_run_font(r, size=11.5, color="555555")
    format_paragraph(p, before=0, after=10, line=1.1)

    add_label_value_table(
        doc,
        [
            ("门店名称", "炖时光"),
            ("门店位置", "意库66栋，暖心素火锅隔壁"),
            ("文档版本", "v1.1"),
            ("适用场景", "线上点单、预订备货、现场现单、会员运营、经营统计"),
            ("营业时间", "11:00-19:00"),
        ],
        widths=(1.7, 4.8),
    )

    add_heading(doc, "1. 项目背景与目标", 1)
    add_body(doc, "当前门店主要依赖群接龙售卖，存在重复发单、人工整理订单、前一晚预订与第二天新单重叠、库存不好把控、复购触达弱等问题。")
    add_body(doc, "本项目目标是把炖时光做成一套可长期经营的小程序系统，让顾客更快下单，让店长更省力备货，并且能通过会员、积分、储值、优惠券和数据分析提升复购。")

    add_heading(doc, "2. 已确认经营规则", 1)
    add_bullet(doc, "营业时间：11:00-19:00。")
    add_bullet(doc, "配送范围：金山谷、保利、意库。")
    add_bullet(doc, "配送费：每单加收 3 元。")
    add_bullet(doc, "已确认价格：红萝卜汤 15 元、凉瓜汤 15 元、冬瓜汤 15 元。")
    add_bullet(doc, "开业活动：全场 88 折。")
    add_bullet(doc, "其他汤品价格：一次性录入商品库，后续店长可在后台直接修改，不需要反复群发。")
    add_bullet(doc, "支付方式：优先微信支付，同时保留线下确认支付作为兜底方式。")

    add_heading(doc, "3. 核心业务流程", 1)
    add_number(doc, "前一晚开放预订，形成预订单。")
    add_number(doc, "到截止时间后自动锁单，进入备货统计。")
    add_number(doc, "第二天早上店长查看采购报表，按预订单和安全余量采购食材。")
    add_number(doc, "店长在后台手动确认库存，必要时可微调。")
    add_number(doc, "11:00 后开放当天现单，只卖剩余库存。")
    add_number(doc, "19:00 自动关单，进入次日预订准备。")

    add_heading(doc, "4. 功能范围", 1)
    add_heading(doc, "4.1 顾客端", 2)
    for item in [
        "首页：营业状态、热销推荐、88 折活动、限时折扣、门店导航、快速下单。",
        "菜单页：汤品、面品、套餐、库存剩余、售罄状态、价格展示。",
        "组合套餐页：先选汤，再选面，系统自动算价，不用备注拼单。",
        "下单页：联系人、手机号、自提/配送、地址、取餐时间、备注、优惠券、积分抵扣。",
        "订单页：待支付、待接单、备餐中、配送中、已完成、历史订单、一键再来一单。",
        "会员中心：积分、优惠券、储值余额、连续下单奖励、复购记录。",
        "通知中心：订单提醒、上新提醒、限时折扣提醒、回流券提醒、积分到账提醒。",
        "门店页：地址、路线图、营业时间、配送范围、联系电话。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.2 店长端", 2)
    for item in [
        "订单管理：接单、改状态、取消、筛选、导出。",
        "商品管理：新增、编辑、上下架、改价、改图、改规格。",
        "库存管理：随时更改库存，查看剩余份数、预订占用、现单可售。",
        "备货报表：第二天早上自动汇总各汤、各面、套餐、配送单、自提单。",
        "营销管理：88 折、限时折扣、连续下单奖励、储值充值活动、优惠券。",
        "数据分析：热销汤品、下单时段、高复购用户、活动效果、客单价、售罄率。",
        "通知管理：订单通知、活动通知、回流通知。",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5. 组合套餐规则", 1)
    add_body(doc, "组合套餐不做备注式拼单，直接做成标准化搭配。顾客先选汤，再选面，系统自动计算总价与折扣后价格。")
    add_bullet(doc, "汤和面分别选择，避免下错单。")
    add_bullet(doc, "系统自动计算原价、折扣价和应付金额。")
    add_bullet(doc, "后台自动统计“汤 + 面”组合销量，方便后续优化推荐。")

    add_heading(doc, "6. 库存与备货逻辑", 1)
    add_body(doc, "库存不是固定死的，而是根据第二天早上采购的实际情况随时可调。因为前一晚已经发出预订，所以系统要把预订单和当天现单分开管理，避免重叠。")
    add_bullet(doc, "每个商品都显示剩余份数。")
    add_bullet(doc, "下单后先锁库存，避免超卖。")
    add_bullet(doc, "店长可随时手动加减库存。")
    add_bullet(doc, "售罄后自动变灰，不再允许下单。")
    add_bullet(doc, "预订单与现单分层展示，避免前后订单混淆。")

    add_heading(doc, "7. 会员、积分与储值活动", 1)
    add_heading(doc, "7.1 积分规则", 2)
    add_bullet(doc, "建议 1 元 = 1 积分。")
    add_bullet(doc, "积分可用于抵扣炖汤或兑换指定商品。")
    add_bullet(doc, "可设置连续下单奖励积分。")
    add_bullet(doc, "建议基础兑换规则：300 分兑换 1 份炖汤。")
    add_heading(doc, "7.2 储值活动", 2)
    add_matrix_table(
        doc,
        ["充值金额", "赠送内容", "说明"],
        [
            ["300 元", "3 张 88 折券", "全单可用，建议设有效期"],
            ["500 元", "6 张 88 折券 + 炖汤 2 份", "适合中度复购用户"],
            ["1000 元", "10 张 88 折券 + 炖汤 5 份", "适合高复购用户"],
        ],
        widths=(1.3, 2.9, 2.3),
    )
    add_body(doc, "充值优惠与优惠券建议分开记录，储值余额和券包分别展示，避免店长核销混乱。")

    add_heading(doc, "8. 促销策略", 1)
    add_bullet(doc, "保留开业全场 88 折。")
    add_bullet(doc, "保留限时折扣。")
    add_bullet(doc, "保留连续下单奖励。")
    add_bullet(doc, "后续如需增强复购，可增加新人首单券、回流券和储值券包，但首版不强制展开更多玩法。")

    add_heading(doc, "9. 数据分析模块", 1)
    add_body(doc, "数据分析是这套系统的重点，目的是帮助店长用数据做采购、定价和促销判断。")
    add_matrix_table(
        doc,
        ["分析项", "统计内容", "用途"],
        [
            ["热销汤品", "哪个汤最受欢迎，按销量/订单数/销售额/复购率统计", "决定备货和主推商品"],
            ["下单时段", "哪个时段下单最多，按小时和午晚市统计", "安排备货和通知时点"],
            ["高复购用户", "30 天内下单次数、消费金额、最近下单时间", "发券和做回流"],
            ["活动效果", "88 折、限时折扣、储值活动核销情况", "判断活动是否有效"],
            ["经营健康度", "客单价、售罄率、取消率、配送占比", "优化门店运营"],
        ],
        widths=(1.2, 3.1, 2.2),
    )

    add_heading(doc, "10. 页面结构", 1)
    add_matrix_table(
        doc,
        ["模块", "页面", "核心内容"],
        [
            ["前台", "首页", "营业状态、活动、快速下单、导航"],
            ["前台", "菜单页", "汤品、面品、套餐、库存、售罄"],
            ["前台", "组合套餐页", "先选汤，再选面，自动计价"],
            ["前台", "下单页", "地址、自提/配送、时间、券、积分"],
            ["前台", "订单页", "订单状态、历史订单、一键复购"],
            ["前台", "会员中心", "积分、券、储值、连续下单奖励"],
            ["后台", "订单管理", "接单、改状态、筛选、导出"],
            ["后台", "商品与库存", "改价、改库存、上下架、售罄"],
            ["后台", "数据分析", "热销、时段、复购、活动、客单价"],
            ["后台", "备货报表", "前一晚预订和次日采购建议"],
        ],
        widths=(1.0, 1.9, 3.6),
    )

    add_heading(doc, "11. 适合单店长的操作原则", 1)
    add_bullet(doc, "后台操作要少，入口要集中。")
    add_bullet(doc, "所有常用动作放在订单、库存、商品、报表四个入口里。")
    add_bullet(doc, "每天早上只需要看报表、改库存、确认开单。")
    add_bullet(doc, "不要把后台做成复杂 ERP，够用、清楚、不会出错最重要。")

    add_heading(doc, "12. 第一版建议范围", 1)
    add_bullet(doc, "顾客下单。")
    add_bullet(doc, "组合套餐自动计价。")
    add_bullet(doc, "库存剩余份数与售罄控制。")
    add_bullet(doc, "预订与现单分流。")
    add_bullet(doc, "会员积分与储值活动。")
    add_bullet(doc, "限时折扣与连续下单奖励。")
    add_bullet(doc, "订单通知与回流提醒。")
    add_bullet(doc, "经营数据分析和备货报表。")

    add_heading(doc, "13. 需要店长最终确认的少量规则", 1)
    add_bullet(doc, "积分是否按 1 元 = 1 分执行。")
    add_bullet(doc, "300 分换 1 份炖汤是否确定为基础规则。")
    add_bullet(doc, "储值券是否有有效期，建议 30 天或 60 天。")
    add_bullet(doc, "预订截止时间建议定在几点。")
    add_bullet(doc, "当天现单从几点开始开放。")
    add_bullet(doc, "88 折是否只用于开业期，还是长期活动。")

    add_heading(doc, "14. 结论", 1)
    add_body(doc, "炖时光适合做成一个“下单 + 备货 + 会员 + 储值 + 数据分析”一体的小程序。对单店长来说，最关键不是功能花哨，而是把预订、现单、库存和报表做顺，让店长一个人也能稳定运营。")

    doc.save(OUTPUT)


if __name__ == "__main__":
    main()
